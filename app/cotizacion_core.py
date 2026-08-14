#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""
cotizador_gto.generar_cotizacion
=================================

Módulo reutilizable para generar cotizaciones de GTO Electrónica y Eléctrico
de Potencia en formato .docx, a partir del formato aprobado con el cliente.

Este archivo es un proyecto NUEVO e independiente: no modifica ni depende de
ningún sistema o archivo previo. Todo lo que necesita (logo, estilos, textos
por defecto) vive dentro de esta misma carpeta.

Uso básico
----------
    from generar_cotizacion import generar_cotizacion, Item

    generar_cotizacion(
        salida_path="Cotizacion_144.docx",
        numero_cotizacion="144-2026",
        fecha_emision="13/08/2026",
        fecha_vencimiento="12/09/2026",
        referencia="SOLPED 0020001200",
        cliente={
            "ruc": "20123456789",
            "razon_social": "MI CLIENTE S.A.C.",
            "direccion": "Av. Siempre Viva 123, Lima",
            "contacto": "Juan Pérez",
            "telefono": "999 999 999",
            "correo": "juan.perez@cliente.com",
        },
        items=[
            Item("Servicio de mantenimiento de equipo X", 1, 2500.00),
        ],
        operacion_gravada=True,   # False si es exportación de servicios (sin IGV)
        activities=[              # opcional: se omite solo si no se pasa nada
            "Actividad 1",
            "Actividad 2",
        ],
    )

Ver ejemplo_uso.py para un caso completo y comentado.
"""

import os
from dataclasses import dataclass, field
from typing import List, Optional, Dict

from docx import Document
from docx.shared import Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_ALIGN_VERTICAL
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

# ============================================================================
# Paleta e identidad de marca (GTO)
# ============================================================================
BRAND_GREEN_HEX = "00944F"
GREEN = RGBColor(0x00, 0x94, 0x4F)
GREEN_DARK = RGBColor(0x00, 0x6B, 0x39)
GREEN_LIGHT_HEX = "E4F3EB"
GRAY_LIGHT_HEX = "F2F2F2"
GRAY_TEXT = RGBColor(0x4A, 0x4A, 0x4A)
WHITE = RGBColor(0xFF, 0xFF, 0xFF)
BORDER_GRAY_HEX = "D9D9D9"
FONT = "Calibri"

TITLE_GAP = 6      # espacio entre la barra verde de sección y su contenido
SECTION_GAP = 12   # espacio entre el final de un bloque y el siguiente

ASSETS_DIR = os.path.join(os.path.dirname(os.path.abspath(__file__)), "assets")
DEFAULT_LOGO_PATH = os.path.join(ASSETS_DIR, "gto.png")

# Datos fijos de la empresa (no cambian de una cotización a otra).
# Se pueden sobrescribir por parámetro si algún día abren otra sede o RUC.
EMPRESA_DEFAULT = {
    "ruc": "20519998786",
    "sede_fiscal": "Cal. German Schreiber Nro. 276 Int. T301 Urb. Santa Ana "
                    "(entre cuadra 1 y 2 de Canaval y Moreyra), San Isidro - Lima",
    "sede_operativa": "Av. Variante Uchumayo Km. 5.5 Lt. 657 Int. A1 Sec. Alto Cural, Arequipa",
    "telefono": "921 502 895 / 932 439 863",
    "correo": "servicios@gtoperu.com",
    "correo_facturacion": "administracion@gtoperu.com",
    "web": "www.gtoperu.com",
    "nombre_footer": "GTO Electrónica y Eléctrico de Potencia",
}

CONDICIONES_DEFAULT = {
    "forma_pago": "Crédito a 30 días – factura negociable.",
    "lugar_entrega": "Por coordinar con el cliente.",
    "plazo_entrega": "Por coordinar con el cliente.",
    "garantia": "6 meses.",
    "validez": "30 días hábiles desde la fecha de emisión.",
    "penalidad": "50% del total del servicio.",
}

CUENTAS_BANCARIAS_DEFAULT = [
    ("Interbank", "Soles", "3003005300733", "00330000300530073319"),
    ("Interbank", "Dólares", "3003005300740", "00330000300530074014"),
    ("Banco de la Nación", "Detracción", "00-101-205436", "-"),
]

GARANTIA_EXCLUSIONES_DEFAULT = [
    "Daños producto de una mala o incorrecta manipulación por parte del cliente.",
    "Daños por incumplimiento de las recomendaciones del fabricante.",
    "Daños producto de transporte o almacenamiento inapropiado.",
    "Daños por falta de cuidado del transportista, manipulación inadecuada o "
    "ausencia de protección ante condiciones climáticas adversas.",
    "Daños por negligencia, actos vandálicos o accidentes.",
    "Signos de manipulación, desarme, cambios o reemplazo de piezas no autorizadas.",
    "Signos de manipulación o ausencia de la placa de control de GTO.",
]


@dataclass
class Item:
    """Una línea de la tabla de ítems."""
    descripcion: str
    cantidad: float
    valor_unitario: float


# ============================================================================
# Número a letras (para la línea "SON:")
# ============================================================================
_UNIDADES = ["", "UNO", "DOS", "TRES", "CUATRO", "CINCO", "SEIS", "SIETE", "OCHO", "NUEVE"]
_DIEZ_A_DIECINUEVE = ["DIEZ", "ONCE", "DOCE", "TRECE", "CATORCE", "QUINCE",
                       "DIECISEIS", "DIECISIETE", "DIECIOCHO", "DIECINUEVE"]
_DECENAS = ["", "", "VEINTE", "TREINTA", "CUARENTA", "CINCUENTA",
            "SESENTA", "SETENTA", "OCHENTA", "NOVENTA"]
_CENTENAS = ["", "CIENTO", "DOSCIENTOS", "TRESCIENTOS", "CUATROCIENTOS",
             "QUINIENTOS", "SEISCIENTOS", "SETECIENTOS", "OCHOCIENTOS", "NOVECIENTOS"]


def _tres_digitos(n: int) -> str:
    if n == 0:
        return ""
    if n == 100:
        return "CIEN"
    c, resto = divmod(n, 100)
    partes = []
    if c:
        partes.append(_CENTENAS[c])
    if resto:
        if resto < 10:
            partes.append(_UNIDADES[resto])
        elif resto < 20:
            partes.append(_DIEZ_A_DIECINUEVE[resto - 10])
        elif resto < 30:
            partes.append("VEINTE" if resto == 20 else "VEINTI" + _UNIDADES[resto - 20])
        else:
            d, u = divmod(resto, 10)
            partes.append(_DECENAS[d] + (" Y " + _UNIDADES[u] if u else ""))
    return " ".join(partes)


def numero_a_letras(monto: float, moneda_letras: str = "DÓLARES AMERICANOS") -> str:
    """Convierte un monto (ej. 2950.00) en su forma escrita para el 'SON:'.
    Verificado contra los ejemplos reales de GTO:
      2950.00 -> "DOS MIL NOVECIENTOS CINCUENTA CON 00/100 DÓLARES AMERICANOS"
       236.00 -> "DOSCIENTOS TREINTA Y SEIS CON 00/100 SOLES"
    """
    entero = int(monto)
    centavos = round((monto - entero) * 100)
    if centavos == 100:
        entero += 1
        centavos = 0

    if entero == 0:
        letras = "CERO"
    else:
        millones, resto = divmod(entero, 1_000_000)
        miles, cientos = divmod(resto, 1000)
        partes = []
        if millones:
            partes.append("UN MILLON" if millones == 1 else f"{_tres_digitos(millones)} MILLONES")
        if miles:
            partes.append("MIL" if miles == 1 else f"{_tres_digitos(miles)} MIL")
        if cientos:
            partes.append(_tres_digitos(cientos))
        letras = " ".join(p for p in partes if p)

    texto = f"{letras} CON {centavos:02d}/100 {moneda_letras}"
    return texto[0] + texto[1:].lower()  # "Dos mil novecientos... con 00/100 dólares..."


# ============================================================================
# Helpers de bajo nivel (OOXML / python-docx)
# ============================================================================

def _set_cell_shading(cell, hex_color):
    tcPr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement('w:shd')
    shd.set(qn('w:val'), 'clear')
    shd.set(qn('w:color'), 'auto')
    shd.set(qn('w:fill'), hex_color)
    tcPr.append(shd)


def _set_cell_borders(cell, color=BORDER_GRAY_HEX, sz=4, sides=("top", "bottom", "left", "right")):
    tcPr = cell._tc.get_or_add_tcPr()
    tcBorders = OxmlElement('w:tcBorders')
    for side in sides:
        el = OxmlElement(f'w:{side}')
        el.set(qn('w:val'), 'single')
        el.set(qn('w:sz'), str(sz))
        el.set(qn('w:space'), '0')
        el.set(qn('w:color'), color)
        tcBorders.append(el)
    tcPr.append(tcBorders)


def _remove_table_borders(table):
    tblPr = table._tbl.tblPr
    borders = OxmlElement('w:tblBorders')
    for side in ('top', 'left', 'bottom', 'right', 'insideH', 'insideV'):
        el = OxmlElement(f'w:{side}')
        el.set(qn('w:val'), 'none')
        el.set(qn('w:sz'), '0')
        el.set(qn('w:space'), '0')
        el.set(qn('w:color'), 'auto')
        borders.append(el)
    tblPr.append(borders)


def _set_col_widths(table, widths_cm):
    table.autofit = False
    for row in table.rows:
        for idx, w in enumerate(widths_cm):
            row.cells[idx].width = Cm(w)
    for idx, w in enumerate(widths_cm):
        table.columns[idx].width = Cm(w)


def _set_cell_margins(cell, top=50, bottom=50, left=100, right=100):
    tcPr = cell._tc.get_or_add_tcPr()
    mar = OxmlElement('w:tcMar')
    for side, val in (('top', top), ('bottom', bottom), ('left', left), ('right', right)):
        node = OxmlElement(f'w:{side}')
        node.set(qn('w:w'), str(val))
        node.set(qn('w:type'), 'dxa')
        mar.append(node)
    tcPr.append(mar)


def _add_run(paragraph, text, bold=False, italic=False, size=10, color=GRAY_TEXT,
             font=FONT, caps=False, spacing=None):
    r = paragraph.add_run(text)
    r.bold = bold
    r.italic = italic
    r.font.size = Pt(size)
    r.font.color.rgb = color
    r.font.name = font
    r.font.all_caps = caps
    if spacing is not None:
        rPr = r._element.get_or_add_rPr()
        spc = OxmlElement('w:spacing')
        spc.set(qn('w:val'), str(spacing))
        rPr.append(spc)
    return r


def _cell_text(cell, text, bold=False, size=9.5, color=GRAY_TEXT, align=WD_ALIGN_PARAGRAPH.LEFT,
               valign=WD_ALIGN_VERTICAL.CENTER, italic=False, caps=False):
    cell.vertical_alignment = valign
    p = cell.paragraphs[0]
    p.alignment = align
    p.paragraph_format.space_after = Pt(0)
    p.paragraph_format.space_before = Pt(0)
    _add_run(p, text, bold=bold, size=size, color=color, italic=italic, caps=caps)
    return p


def _spacer(doc, pts=8):
    """Espacio en blanco de altura EXACTA (ritmo vertical consistente)."""
    p = doc.add_paragraph()
    pf = p.paragraph_format
    pf.space_before = Pt(0)
    pf.space_after = Pt(0)
    pf.line_spacing_rule = WD_LINE_SPACING.EXACTLY
    pf.line_spacing = Pt(pts)
    return p


def _bullet(doc, text, bold_prefix=None, size=9.5, color=GRAY_TEXT):
    p = doc.add_paragraph(style='List Bullet')
    p.paragraph_format.space_after = Pt(3)
    p.paragraph_format.left_indent = Cm(0.5)
    if bold_prefix:
        _add_run(p, bold_prefix, bold=True, size=size, color=RGBColor(0x2A, 0x2A, 0x2A))
    _add_run(p, text, size=size, color=color)
    return p


def _section_bar(doc, title):
    table = doc.add_table(rows=1, cols=1)
    _remove_table_borders(table)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    _set_col_widths(table, [18.0])
    cell = table.cell(0, 0)
    _set_cell_shading(cell, BRAND_GREEN_HEX)
    p = cell.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    p.paragraph_format.space_before = Pt(3)
    p.paragraph_format.space_after = Pt(3)
    _add_run(p, title, bold=True, size=10.5, color=WHITE, caps=True, spacing=15)
    _set_cell_margins(cell, top=60, bottom=60, left=150, right=150)
    _spacer(doc, TITLE_GAP)
    return table


def _add_page_number_field(paragraph):
    run = paragraph.add_run()
    fld1 = OxmlElement('w:fldChar'); fld1.set(qn('w:fldCharType'), 'begin')
    instr = OxmlElement('w:instrText'); instr.set(qn('xml:space'), 'preserve'); instr.text = 'PAGE'
    fld2 = OxmlElement('w:fldChar'); fld2.set(qn('w:fldCharType'), 'end')
    run._r.append(fld1); run._r.append(instr); run._r.append(fld2)


def _add_numpages_field(paragraph):
    run = paragraph.add_run()
    fld1 = OxmlElement('w:fldChar'); fld1.set(qn('w:fldCharType'), 'begin')
    instr = OxmlElement('w:instrText'); instr.set(qn('xml:space'), 'preserve'); instr.text = 'NUMPAGES'
    fld2 = OxmlElement('w:fldChar'); fld2.set(qn('w:fldCharType'), 'end')
    run._r.append(fld1); run._r.append(instr); run._r.append(fld2)


# ============================================================================
# Función principal
# ============================================================================

def generar_cotizacion(
    salida_path: str,
    numero_cotizacion: str,
    fecha_emision: str,
    fecha_vencimiento: str,
    cliente: Dict[str, str],
    items: List[Item],
    referencia: str = "",
    asesor: Optional[Dict[str, str]] = None,
    moneda_simbolo: str = "US$",
    moneda_letras: str = "DÓLARES AMERICANOS",
    igv_pct: int = 18,
    operacion_gravada: bool = True,
    activities: Optional[List[str]] = None,
    condiciones: Optional[Dict[str, str]] = None,
    cuentas_bancarias: Optional[List[tuple]] = None,
    garantia_texto: Optional[str] = None,
    garantia_exclusiones: Optional[List[str]] = None,
    empresa: Optional[Dict[str, str]] = None,
    logo_path: Optional[str] = None,
) -> str:
    """Genera una cotización .docx con el formato aprobado de GTO.

    Parámetros clave para automatización
    -------------------------------------
    operacion_gravada : bool
        True  -> venta/servicio LOCAL: se muestra "Operación gravada" + IGV + total.
        False -> venta/servicio EXTERNO (exportación de servicios, cliente/uso
                 fuera del Perú, Art. 33° Ley del IGV): se omite la fila de IGV
                 y se agrega automáticamente la nota legal correspondiente.
    activities : list[str] | None
        Lista de actividades del "Alcance del servicio". Si es None o [], la
        sección completa se omite sin dejar espacio en blanco.
    condiciones : dict | None
        Sobrescribe cualquiera de las llaves de CONDICIONES_DEFAULT
        (forma_pago, lugar_entrega, plazo_entrega, garantia, validez, penalidad).

    Devuelve la ruta del archivo generado.
    """
    asesor = asesor or {
        "nombre": "Juan Alexis Gómez Mamani",
        "celular": "921 502 895",
        "correo": "servicios@gtoperu.com",
    }
    cond = {**CONDICIONES_DEFAULT, **(condiciones or {})}
    cuentas_bancarias = cuentas_bancarias or CUENTAS_BANCARIAS_DEFAULT
    garantia_exclusiones = garantia_exclusiones or GARANTIA_EXCLUSIONES_DEFAULT
    empresa = {**EMPRESA_DEFAULT, **(empresa or {})}
    logo_path = logo_path or DEFAULT_LOGO_PATH

    def fmt_money(v):
        return f"{moneda_simbolo} {v:,.2f}"

    doc = Document()

    section = doc.sections[0]
    section.page_width = Cm(21.0)
    section.page_height = Cm(29.7)
    section.top_margin = Cm(1.3)
    section.bottom_margin = Cm(1.3)
    section.left_margin = Cm(1.5)
    section.right_margin = Cm(1.5)

    normal = doc.styles['Normal']
    normal.font.name = FONT
    normal.font.size = Pt(9.5)
    normal.font.color.rgb = GRAY_TEXT
    normal.paragraph_format.space_after = Pt(0)

    # ---------------- Encabezado ----------------
    header_tbl = doc.add_table(rows=1, cols=2)
    _remove_table_borders(header_tbl)
    _set_col_widths(header_tbl, [12.0, 6.0])

    left_cell = header_tbl.cell(0, 0)
    left_cell.vertical_alignment = WD_ALIGN_VERTICAL.TOP
    p = left_cell.paragraphs[0]
    p.paragraph_format.space_after = Pt(6)
    if logo_path and os.path.exists(logo_path):
        p.add_run().add_picture(logo_path, width=Cm(4.6))

    address_lines = [
        ("Sede fiscal: ", empresa["sede_fiscal"]),
        ("Sede operativa: ", empresa["sede_operativa"]),
    ]
    for label, addr in address_lines:
        pp = left_cell.add_paragraph()
        pp.paragraph_format.space_after = Pt(1)
        _add_run(pp, label, bold=True, size=8.5, color=GRAY_TEXT)
        _add_run(pp, addr, size=8.5, color=GRAY_TEXT)

    for line in [f"Tel: {empresa['telefono']}   |   {empresa['correo']}", empresa["web"]]:
        pp = left_cell.add_paragraph()
        pp.paragraph_format.space_after = Pt(1)
        _add_run(pp, line, size=8.5, color=GRAY_TEXT)

    right_cell = header_tbl.cell(0, 1)
    right_cell.vertical_alignment = WD_ALIGN_VERTICAL.TOP
    rp = right_cell.paragraphs[0]
    rp.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    rp.paragraph_format.space_after = Pt(2)
    _add_run(rp, "COTIZACIÓN", bold=True, size=20, color=GREEN, caps=True, spacing=20)

    rp2 = right_cell.add_paragraph()
    rp2.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    rp2.paragraph_format.space_after = Pt(2)
    _add_run(rp2, f"N° {numero_cotizacion}", bold=True, size=12.5, color=GREEN_DARK)

    rp3 = right_cell.add_paragraph()
    rp3.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    _add_run(rp3, f"R.U.C. {empresa['ruc']}", size=9, color=GRAY_TEXT)

    sep = doc.add_paragraph()
    sep.paragraph_format.space_before = Pt(0)
    sep.paragraph_format.space_after = Pt(8)
    pPr = sep._p.get_or_add_pPr()
    pBdr = OxmlElement('w:pBdr')
    bottom = OxmlElement('w:bottom')
    bottom.set(qn('w:val'), 'single')
    bottom.set(qn('w:sz'), '18')
    bottom.set(qn('w:space'), '1')
    bottom.set(qn('w:color'), BRAND_GREEN_HEX)
    pBdr.append(bottom)
    pPr.append(pBdr)

    # ---------------- Datos cliente / cotización ----------------
    info_rows = [
        ("R.U.C. / DNI", cliente.get("ruc", "-"), "N° de cotización", numero_cotizacion),
        ("Razón social", cliente.get("razon_social", "-"), "Fecha de emisión", fecha_emision),
        ("Dirección", cliente.get("direccion", "-"), "Fecha de vencimiento", fecha_vencimiento),
        ("Contacto", cliente.get("contacto", "-"), "Referencia", referencia or "-"),
        ("Teléfono", cliente.get("telefono", "-"), "Asesor comercial", asesor["nombre"]),
        ("Correo", cliente.get("correo", "-"), "Celular / correo asesor",
         f"{asesor['celular']} / {asesor['correo']}"),
    ]

    info_tbl = doc.add_table(rows=len(info_rows), cols=4)
    info_tbl.alignment = WD_TABLE_ALIGNMENT.CENTER
    _set_col_widths(info_tbl, [3.6, 5.2, 4.0, 5.2])
    for i, (l1, v1, l2, v2) in enumerate(info_rows):
        row = info_tbl.rows[i]
        for c in row.cells:
            _set_cell_borders(c, color=BORDER_GRAY_HEX, sz=4)
            _set_cell_margins(c, top=45, bottom=45, left=110, right=110)
        _set_cell_shading(row.cells[0], GRAY_LIGHT_HEX)
        _set_cell_shading(row.cells[2], GRAY_LIGHT_HEX)
        _cell_text(row.cells[0], l1, bold=True, size=9, color=RGBColor(0x33, 0x33, 0x33))
        _cell_text(row.cells[1], v1, size=9)
        _cell_text(row.cells[2], l2, bold=True, size=9, color=RGBColor(0x33, 0x33, 0x33))
        _cell_text(row.cells[3], v2, size=9)

    _spacer(doc, SECTION_GAP)

    # ---------------- Tabla de ítems ----------------
    items_tbl = doc.add_table(rows=1 + len(items), cols=5)
    items_tbl.alignment = WD_TABLE_ALIGNMENT.CENTER
    _set_col_widths(items_tbl, [1.4, 8.6, 1.6, 3.1, 3.3])

    headers = ["ÍTEM", "DESCRIPCIÓN", "CANT.", f"V. UNIT. ({moneda_simbolo})", f"IMPORTE ({moneda_simbolo})"]
    for idx, h in enumerate(headers):
        c = items_tbl.rows[0].cells[idx]
        _set_cell_shading(c, BRAND_GREEN_HEX)
        _set_cell_margins(c, top=70, bottom=70, left=100, right=100)
        align = WD_ALIGN_PARAGRAPH.CENTER if idx != 1 else WD_ALIGN_PARAGRAPH.LEFT
        _cell_text(c, h, bold=True, size=9, color=WHITE, align=align, caps=True)

    subtotal_val = 0.0
    for r_i, item in enumerate(items, start=1):
        importe = item.cantidad * item.valor_unitario
        subtotal_val += importe
        row = items_tbl.rows[r_i]
        fill = "FFFFFF" if r_i % 2 else GRAY_LIGHT_HEX
        for c in row.cells:
            _set_cell_shading(c, fill)
            _set_cell_borders(c, color=BORDER_GRAY_HEX, sz=4)
            _set_cell_margins(c, top=70, bottom=70, left=100, right=100)
        _cell_text(row.cells[0], str(r_i), align=WD_ALIGN_PARAGRAPH.CENTER, size=9)
        dcell = row.cells[1]
        dcell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
        dp = dcell.paragraphs[0]
        dp.paragraph_format.space_after = Pt(0)
        _add_run(dp, item.descripcion, size=9)
        _cell_text(row.cells[2], f"{item.cantidad:g}", align=WD_ALIGN_PARAGRAPH.CENTER, size=9)
        _cell_text(row.cells[3], f"{item.valor_unitario:,.2f}", align=WD_ALIGN_PARAGRAPH.RIGHT, size=9)
        _cell_text(row.cells[4], f"{importe:,.2f}", align=WD_ALIGN_PARAGRAPH.RIGHT, size=9, bold=True)

    _spacer(doc, SECTION_GAP)

    # ---------------- Totales ----------------
    igv_val = round(subtotal_val * igv_pct / 100, 2) if operacion_gravada else 0.0
    total_val = subtotal_val + igv_val

    tot_outer = doc.add_table(rows=1, cols=2)
    _remove_table_borders(tot_outer)
    _set_col_widths(tot_outer, [10.8, 7.2])
    tot_outer.rows[0].cells[0].text = ""
    tot_cell = tot_outer.rows[0].cells[1]

    if operacion_gravada:
        totals = [
            ("Operación gravada", fmt_money(subtotal_val), False),
            (f"I.G.V. ({igv_pct}%)", fmt_money(igv_val), False),
            ("IMPORTE TOTAL", fmt_money(total_val), True),
        ]
    else:
        totals = [
            ("Operación exonerada", fmt_money(subtotal_val), False),
            ("IMPORTE TOTAL", fmt_money(total_val), True),
        ]

    tot_tbl_holder = tot_cell.add_table(rows=len(totals), cols=2)
    _set_col_widths(tot_tbl_holder, [4.0, 3.2])
    _remove_table_borders(tot_tbl_holder)
    for i, (label, val, strong) in enumerate(totals):
        row = tot_tbl_holder.rows[i]
        for c in row.cells:
            _set_cell_margins(c, top=45, bottom=45, left=90, right=90)
        bg = BRAND_GREEN_HEX if strong else GRAY_LIGHT_HEX
        fg = WHITE if strong else GRAY_TEXT
        _set_cell_shading(row.cells[0], bg)
        _set_cell_shading(row.cells[1], bg)
        _cell_text(row.cells[0], label, bold=strong, size=10 if strong else 9.5, color=fg, caps=strong)
        _cell_text(row.cells[1], val, bold=strong, size=10 if strong else 9.5, color=fg,
                   align=WD_ALIGN_PARAGRAPH.RIGHT)

    son_p = doc.add_paragraph()
    son_p.paragraph_format.space_before = Pt(6)
    son_p.paragraph_format.space_after = Pt(2)
    _add_run(son_p, "SON: ", bold=True, size=9, color=RGBColor(0x33, 0x33, 0x33))
    _add_run(son_p, numero_a_letras(total_val, moneda_letras) + ".", italic=True, size=9)

    if not operacion_gravada:
        igv_note_p = doc.add_paragraph()
        igv_note_p.paragraph_format.space_after = Pt(2)
        _add_run(igv_note_p, "Operación no gravada con IGV — exportación de servicios "
                              "(Art. 33° de la Ley del IGV).", italic=True, size=8, color=GRAY_TEXT)

    _spacer(doc, SECTION_GAP)

    # ---------------- Alcance del servicio (condicional) ----------------
    if activities:
        _section_bar(doc, "Alcance del servicio")
        for act in activities:
            _bullet(doc, act, size=9.5)
        _spacer(doc, SECTION_GAP)

    # ---------------- Condiciones comerciales ----------------
    _section_bar(doc, "Condiciones comerciales")
    _bullet(doc, cond["forma_pago"], "Forma de pago: ")
    _bullet(doc, cond["lugar_entrega"], "Lugar de recojo y entrega: ")
    _bullet(doc, cond["plazo_entrega"], "Plazo de entrega: ")
    _bullet(doc, cond["garantia"], "Garantía: ")
    _bullet(doc, cond["validez"], "Validez de la cotización: ")
    _bullet(doc, cond["penalidad"], "Penalidad por cancelación de orden de compra: ")
    _bullet(doc, f"{moneda_letras.capitalize()} ({moneda_simbolo}).", "Moneda: ")
    _spacer(doc, SECTION_GAP)

    # ---------------- Política de garantía ----------------
    _section_bar(doc, "Política de garantía")
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(5)
    _add_run(p, garantia_texto or (
        "GTO garantiza sus productos y servicios de acuerdo con el período establecido en la sección de "
        "condiciones comerciales de la presente cotización. Para hacer efectiva la garantía, el cliente "
        "deberá comunicar el evento de manera inmediata al correo " + empresa["correo"] + ", adjuntando un "
        "informe técnico que sustente la falla reportada y la evidencia fotográfica correspondiente."
    ), size=9)
    p2 = doc.add_paragraph()
    p2.paragraph_format.space_after = Pt(5)
    _add_run(p2, "GTO contará con un plazo de 7 a 15 días calendario para evaluar el reclamo de garantía y "
                 "emitir la aceptación o rechazo correspondiente mediante informe técnico.", size=9)
    p3 = doc.add_paragraph()
    p3.paragraph_format.space_after = Pt(3)
    _add_run(p3, "La garantía no aplica en los siguientes casos:", bold=True, size=9,
              color=RGBColor(0x33, 0x33, 0x33))
    for e in garantia_exclusiones:
        _bullet(doc, e, size=8.75)
    p4 = doc.add_paragraph()
    p4.paragraph_format.space_before = Pt(4)
    p4.paragraph_format.space_after = Pt(2)
    _add_run(p4, "GTO no se hace responsable por lucro cesante derivado de la parada del equipo donde se "
                 "instaló el componente.", italic=True, size=8.75)
    _spacer(doc, SECTION_GAP)

    # ---------------- Comunicación ----------------
    _section_bar(doc, "Comunicación")
    _bullet(doc, empresa["correo"], "Consultas técnicas y comerciales: ")
    _bullet(doc, empresa["correo_facturacion"], "Facturación y cambios: ")
    _bullet(doc, "GTO se compromete a informar al cliente cualquier situación que afecte el plazo de "
                 "entrega del servicio.")
    _spacer(doc, SECTION_GAP)

    # ---------------- Medios de pago ----------------
    _section_bar(doc, "Medios de pago")
    pay_tbl = doc.add_table(rows=1 + len(cuentas_bancarias), cols=4)
    pay_tbl.alignment = WD_TABLE_ALIGNMENT.CENTER
    _set_col_widths(pay_tbl, [3.5, 3.5, 5.5, 5.5])
    for idx, h in enumerate(["Banco", "Moneda", "Cuenta corriente", "Cuenta CCI"]):
        c = pay_tbl.rows[0].cells[idx]
        _set_cell_shading(c, "D8E9DE")
        _set_cell_margins(c, top=50, bottom=50, left=100, right=100)
        _cell_text(c, h, bold=True, size=9, color=GREEN_DARK, align=WD_ALIGN_PARAGRAPH.CENTER)
    for i, (b, m, cc, cci) in enumerate(cuentas_bancarias, start=1):
        row = pay_tbl.rows[i]
        fill = "FFFFFF" if i % 2 else GRAY_LIGHT_HEX
        for c in row.cells:
            _set_cell_shading(c, fill)
            _set_cell_borders(c, color=BORDER_GRAY_HEX, sz=4)
            _set_cell_margins(c, top=50, bottom=50, left=100, right=100)
        _cell_text(row.cells[0], b, size=9)
        _cell_text(row.cells[1], m, size=9)
        _cell_text(row.cells[2], cc, size=9)
        _cell_text(row.cells[3], cci, size=9)

    _spacer(doc, SECTION_GAP)
    warn_tbl = doc.add_table(rows=1, cols=1)
    _remove_table_borders(warn_tbl)
    _set_col_widths(warn_tbl, [18.0])
    wcell = warn_tbl.cell(0, 0)
    _set_cell_shading(wcell, GREEN_LIGHT_HEX)
    _set_cell_margins(wcell, top=80, bottom=80, left=180, right=150)
    wBorders = OxmlElement('w:tcBorders')
    wleft = OxmlElement('w:left')
    wleft.set(qn('w:val'), 'single')
    wleft.set(qn('w:sz'), '24')
    wleft.set(qn('w:space'), '0')
    wleft.set(qn('w:color'), BRAND_GREEN_HEX)
    wBorders.append(wleft)
    wcell._tc.get_or_add_tcPr().append(wBorders)
    wp = wcell.paragraphs[0]
    _add_run(wp, "Antes de realizar cualquier pago, confirme que la cuenta y la moneda correspondan al "
                 "comprobante emitido por GTO. GTO no se hace responsable por depósitos realizados a "
                 "cuentas no verificadas.", bold=True, size=8.5, color=GREEN_DARK)

    # ---------------- Pie de página ----------------
    footer = section.footer
    footer.is_linked_to_previous = False
    fp = footer.paragraphs[0]
    fp.alignment = WD_ALIGN_PARAGRAPH.CENTER
    pPr = fp._p.get_or_add_pPr()
    pBdr2 = OxmlElement('w:pBdr')
    top = OxmlElement('w:top')
    top.set(qn('w:val'), 'single'); top.set(qn('w:sz'), '4'); top.set(qn('w:space'), '4')
    top.set(qn('w:color'), 'BFBFBF')
    pBdr2.append(top)
    pPr.append(pBdr2)
    _add_run(fp, f"{empresa['nombre_footer']}   |   R.U.C. {empresa['ruc']}   |   {empresa['web']}   |   Página ",
              size=7.5, color=RGBColor(0x80, 0x80, 0x80))
    _add_page_number_field(fp)
    _add_run(fp, " de ", size=7.5, color=RGBColor(0x80, 0x80, 0x80))
    _add_numpages_field(fp)

    doc.save(salida_path)
    return salida_path
